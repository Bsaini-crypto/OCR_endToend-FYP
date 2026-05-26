"""
Local file storage module for saving OCR records as Word documents.
"""
import json
import os
from datetime import datetime
from typing import Dict, Any, Optional
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class LocalStorage:
    """Handles saving records to local folder."""
    
    def __init__(self, storage_dir: str = "records"):
        """
        Initialize local storage.
        
        Args:
            storage_dir: Directory path where records will be saved
        """
        self.storage_dir = storage_dir
        self._ensure_directory_exists()
    
    def _ensure_directory_exists(self):
        """Create storage directory if it doesn't exist."""
        try:
            os.makedirs(self.storage_dir, exist_ok=True)
            logger.info(f"Storage directory ready: {self.storage_dir}")
        except Exception as e:
            logger.error(f"Failed to create storage directory: {e}")
            raise
    
    def save_record(self, data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Save a record to local storage as Word document (.docx).
        
        Args:
            data: Dictionary containing the record data
            filename: Optional custom filename. If not provided, generates timestamp-based name
            
        Returns:
            Path to the saved file
        """
        try:
            if filename is None:
                # Generate filename with timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                filename = f"record_{timestamp}.docx"
            
            # Ensure .docx extension
            if not filename.endswith('.docx'):
                if filename.endswith('.doc'):
                    filename = filename[:-4] + '.docx'
                else:
                    filename += '.docx'
            
            filepath = os.path.join(self.storage_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title
            title = doc.add_heading('OCR Extracted Data', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp_para = doc.add_paragraph(f"Extracted on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_format = timestamp_para.runs[0].font
            timestamp_format.size = Pt(10)
            timestamp_format.italic = True
            
            # Add a line break
            doc.add_paragraph()
            
            # Recursively add data to document
            self._add_data_to_document(doc, data)
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"Record saved successfully as Word document: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to save record: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise
    
    def _add_data_to_document(self, doc: Document, data: Any, level: int = 0, parent_key: str = ""):
        """
        Recursively add data to Word document with proper formatting.
        
        Args:
            doc: Word document object
            data: Data to add (dict, list, or primitive)
            level: Current nesting level (for indentation)
            parent_key: Parent key name for context
        """
        if isinstance(data, dict):
            for key, value in data.items():
                # Format key name
                key_name = str(key)
                
                # Add heading or paragraph based on level
                if level == 0:
                    # Top-level keys get headings
                    heading = doc.add_heading(key_name, level=1)
                elif level == 1:
                    # Second level gets subheadings
                    heading = doc.add_heading(key_name, level=2)
                else:
                    # Deeper levels get bold paragraphs
                    para = doc.add_paragraph()
                    run = para.add_run(key_name + ":")
                    run.bold = True
                    run.font.size = Pt(11)
                
                # Recursively add value
                if isinstance(value, (dict, list)):
                    self._add_data_to_document(doc, value, level + 1, key_name)
                else:
                    # Add value as paragraph
                    if level >= 2:
                        # Value continues on same line or new line
                        para = doc.add_paragraph(str(value))
                        para.style = 'List Bullet'
                    else:
                        para = doc.add_paragraph(str(value))
                        para_format = para.runs[0].font if para.runs else None
                        if para_format:
                            para_format.size = Pt(11)
                
                # Add spacing between sections
                if level <= 1:
                    doc.add_paragraph()
        
        elif isinstance(data, list):
            # Handle lists
            for i, item in enumerate(data, 1):
                if isinstance(item, dict):
                    # List of objects - add as numbered items
                    para = doc.add_paragraph(f"Item {i}:", style='List Number')
                    para.runs[0].bold = True
                    self._add_data_to_document(doc, item, level + 1, f"Item {i}")
                else:
                    # Simple list items
                    para = doc.add_paragraph(f"{i}. {str(item)}", style='List Bullet')
        
        else:
            # Primitive value - should have been handled above, but just in case
            para = doc.add_paragraph(str(data))
            para_format = para.runs[0].font if para.runs else None
            if para_format:
                para_format.size = Pt(11)

